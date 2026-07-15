from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document

MODULE_ROOT = Path(__file__).resolve().parents[1]
MARKERS_PATH = MODULE_ROOT / "resources" / "markers.json"

def _cp1251_mojibake(text: str) -> str:
    return text.encode("utf-8").decode("cp1251")


BROKEN_MARKER_PREFIXES = {
    _cp1251_mojibake("Анотація"): "\u0410\u043d\u043e\u0442\u0430\u0446\u0456\u044f",
    _cp1251_mojibake("Ключові слова"): "\u041a\u043b\u044e\u0447\u043e\u0432\u0456 \u0441\u043b\u043e\u0432\u0430",
    _cp1251_mojibake("Список використаних джерел"): "\u0421\u043f\u0438\u0441\u043e\u043a \u0432\u0438\u043a\u043e\u0440\u0438\u0441\u0442\u0430\u043d\u0438\u0445 \u0434\u0436\u0435\u0440\u0435\u043b",
    _cp1251_mojibake("Література"): "\u041b\u0456\u0442\u0435\u0440\u0430\u0442\u0443\u0440\u0430",
}

DEFAULT_ANNOTATION_MARKERS = [
    "\u0410\u043d\u043e\u0442\u0430\u0446\u0456\u044f",
    "\u0410\u043d\u043e\u0442\u0430\u0446\u0456\u044f:",
    "\u0410\u043d\u043e\u0442\u0430\u0446\u0456\u044f.",
    "Annotation",
    "Annotation:",
    "Abstract",
    "Abstract:",
]

DEFAULT_KEYWORD_MARKERS = [
    "\u041a\u043b\u044e\u0447\u043e\u0432\u0456 \u0441\u043b\u043e\u0432\u0430",
    "\u041a\u043b\u044e\u0447\u043e\u0432\u0456 \u0441\u043b\u043e\u0432\u0430:",
    "Keywords",
    "Keywords:",
    "Key words",
    "Key Words",
    "Key words:",
    "Key Words:",
]

DEFAULT_REFERENCE_MARKERS = [
    "\u0421\u043f\u0438\u0441\u043e\u043a \u0432\u0438\u043a\u043e\u0440\u0438\u0441\u0442\u0430\u043d\u0438\u0445 \u0434\u0436\u0435\u0440\u0435\u043b",
    "\u0421\u043f\u0438\u0441\u043e\u043a \u0432\u0438\u043a\u043e\u0440\u0438\u0441\u0442\u0430\u043d\u0438\u0445 \u0434\u0436\u0435\u0440\u0435\u043b:",
    "\u041b\u0456\u0442\u0435\u0440\u0430\u0442\u0443\u0440\u0430",
    "References",
    "References:",
    "List of references",
    "List of reference",
]


def _load_list(key: str, fallback: list[str]) -> list[str]:
    try:
        data = json.loads(MARKERS_PATH.read_text(encoding="utf-8"))
    except Exception:
        return fallback
    if not isinstance(data, dict):
        return fallback
    values = data.get(key)
    if not isinstance(values, list):
        return fallback
    cleaned = [str(x).strip() for x in values if str(x).strip()]
    return cleaned or fallback


def _normalize_source(text: str) -> str:
    text = (text or "").replace("\u00a0", " ").lstrip()
    text = re.sub(r"^[Cc](?=писок\b)", "\u0421", text)
    for broken, fixed in BROKEN_MARKER_PREFIXES.items():
        if text.startswith(broken):
            text = fixed + text[len(broken):]
            break
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"^\s*[IVXLCDM]+\s*[\.\)\-]?\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^\s*\d+\s*[\.\)\-]?\s*", "", text)
    return text


def _extract_marker_remainder(text: str, markers: list[str]) -> tuple[str, str | None]:
    source = _normalize_source(text)
    lowered = source.lower()
    cleaned = [m.strip() for m in markers if isinstance(m, str) and m.strip()]
    cleaned.sort(key=len, reverse=True)
    for marker in cleaned:
        m_low = marker.lower()
        if lowered.startswith(m_low):
            remainder = source[len(marker):]
            remainder = re.sub(r"^[\s\.\:\;\-\u2013\u2014]+", "", remainder)
            return remainder, marker
        marker_base = re.sub(r"[\s\.\:\;\-\u2013\u2014]+$", "", marker)
        if marker_base and lowered.startswith(marker_base.lower()):
            remainder = source[len(marker_base):]
            remainder = re.sub(r"^[\s\.\:\;\-\u2013\u2014]+", "", remainder)
            return remainder, marker
    return source, None


def _is_cyrillic(text: str) -> bool:
    return bool(re.search(r"[\u0400-\u04FF]", text or ""))


def _looks_like_other_marker(text: str, other_markers: list[str]) -> bool:
    low = _normalize_source(text).lower()
    for marker in other_markers:
        if marker and low.startswith(marker.lower()):
            return True
    return False


def _pull_next_content(doc: Document, start_idx: int, stop_markers: list[str]) -> tuple[str, int | None]:
    for j in range(start_idx + 1, len(doc.paragraphs)):
        t = (doc.paragraphs[j].text or "").strip()
        if not t:
            continue
        if _looks_like_other_marker(t, stop_markers):
            return "", None
        return t, j
    return "", None


def _capitalize_first(text: str) -> str:
    text = text.lstrip()
    if not text:
        return text
    first = text[0]
    if first.isalpha():
        return first.upper() + text[1:]
    return text


def _lowercase_first(text: str) -> str:
    text = text.lstrip()
    if not text:
        return text
    first = text[0]
    if first.isalpha():
        return first.lower() + text[1:]
    return text


def _replace_paragraph_with_heading(para, heading: str, content: str) -> None:
    for run in list(para.runs):
        keep_run = False
        try:
            for node in run._element.iter():
                tag = str(node.tag)
                if tag.endswith("}drawing") or tag.endswith("}pict") or tag.endswith("}object"):
                    keep_run = True
                    break
        except Exception:
            keep_run = False
        if keep_run:
            continue
        run.text = ""
    heading_run = para.add_run(heading)
    heading_run.bold = True
    if content:
        para.add_run(" ")
        para.add_run(content)


def _keyword_heading(text: str) -> str:
    return "\u041a\u043b\u044e\u0447\u043e\u0432\u0456 \u0441\u043b\u043e\u0432\u0430:" if _is_cyrillic(text or "") else "Keywords:"


def _remove_paragraph(para) -> None:
    try:
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None
    except Exception:
        pass


def _paragraph_has_graphics(para) -> bool:
    try:
        for node in para._element.iter():
            tag = str(node.tag)
            if tag.endswith("}drawing") or tag.endswith("}pict") or tag.endswith("}object"):
                return True
    except Exception:
        return False
    return False


def _is_visually_empty_paragraph(para) -> bool:
    return not (para.text or "").strip() and not _paragraph_has_graphics(para)


def _ensure_single_blank_after(doc: Document, idx: int) -> bool:
    if idx < 0 or idx >= len(doc.paragraphs):
        return False

    changed = False

    if idx + 1 >= len(doc.paragraphs):
        doc.add_paragraph("")
        return True

    next_para = doc.paragraphs[idx + 1]

    if not _is_visually_empty_paragraph(next_para):
        doc.paragraphs[idx + 1].insert_paragraph_before("")
        return True

    j = idx + 2
    while j < len(doc.paragraphs):
        if not _is_visually_empty_paragraph(doc.paragraphs[j]):
            break
        _remove_paragraph(doc.paragraphs[j])
        changed = True

    return changed


def normalize_annotations(
    draft_path: Path,
    output_path: Path | None = None,
    *,
    logs_dir: Path | None = None,
    write_logs: bool = True,
) -> Path:
    markers = _load_list("annotation_markers", DEFAULT_ANNOTATION_MARKERS)
    keyword_markers = _load_list("keyword_markers", DEFAULT_KEYWORD_MARKERS)
    reference_markers = _load_list("reference_markers", DEFAULT_REFERENCE_MARKERS)
    stop_markers = keyword_markers + reference_markers + markers

    doc = Document(draft_path)
    changes = []
    keyword_spacing_changes = []
    keyword_changes = []
    found = 0
    pulled = 0
    keyword_pulled = 0

    for idx, para in enumerate(doc.paragraphs):
        text = para.text or ""
        remainder, marker = _extract_marker_remainder(text, markers)
        if marker is None:
            continue

        found += 1
        lang = "uk" if _is_cyrillic(marker or text) else "en"
        heading = "\u0410\u043d\u043e\u0442\u0430\u0446\u0456\u044f." if lang == "uk" else "Abstract."

        content = remainder.strip()
        pulled_from = None

        if not content:
            content, pulled_from = _pull_next_content(doc, idx, stop_markers)
            if pulled_from is not None:
                doc.paragraphs[pulled_from].text = ""
                pulled += 1

        content = _capitalize_first(content)
        new_text = f"{heading} {content}" if content else heading

        if (para.text or "") != new_text:
            changes.append(
                {
                    "index": idx + 1,
                    "before": para.text,
                    "after": new_text,
                    "pulled_from": (pulled_from + 1) if pulled_from is not None else None,
                }
            )
            _replace_paragraph_with_heading(para, heading, content)

    annotation_hits = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text or ""
        _, marker = _extract_marker_remainder(text, markers)
        if marker is not None:
            annotation_hits.append(idx)

    for idx in reversed(annotation_hits):
        j = idx + 1
        while j < len(doc.paragraphs):
            if not _is_visually_empty_paragraph(doc.paragraphs[j]):
                break
            _remove_paragraph(doc.paragraphs[j])

    keyword_hits = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text or ""
        remainder, marker = _extract_marker_remainder(text, keyword_markers)
        if marker is not None:
            keyword_hits.append(idx)
            heading = _keyword_heading(marker or text)
            content = remainder.strip()
            pulled_from = None

            if not content:
                content, pulled_from = _pull_next_content(doc, idx, stop_markers)
                if pulled_from is not None:
                    doc.paragraphs[pulled_from].text = ""
                    keyword_pulled += 1

            content = _lowercase_first(content)
            new_text = f"{heading} {content}" if content else heading
            before_text = para.text or ""

            # Always rewrite keyword heading to enforce bold marker style.
            _replace_paragraph_with_heading(para, heading, content)

            if before_text != new_text:
                keyword_changes.append(
                    {
                        "index": idx + 1,
                        "before": before_text,
                        "after": new_text,
                        "pulled_from": (pulled_from + 1) if pulled_from is not None else None,
                    }
                )
                # Already rewritten above to enforce style even when text is unchanged.

    for idx in reversed(keyword_hits):
        if _ensure_single_blank_after(doc, idx):
            keyword_spacing_changes.append(idx + 1)

    target = output_path or draft_path.with_name(draft_path.stem + "_annotation.docx")
    doc.save(target)

    if not write_logs:
        return target

    report = {
        "input": str(draft_path),
        "output": str(target),
        "markers_found": found,
        "changed": len(changes),
        "pulled_from_next": pulled,
        "keywords_changed": len(keyword_changes),
        "keywords_pulled_from_next": keyword_pulled,
        "keyword_spacing_fixed": len(keyword_spacing_changes),
        "changes": changes,
        "keyword_changes": keyword_changes,
    }

    if logs_dir is None:
        logs_dir = target.parent

    logs_dir.mkdir(parents=True, exist_ok=True)

    (logs_dir / f"{target.stem}.annotation_report.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    lines = [
        f"[{item['index']}] pulled_from={item['pulled_from']} | {item['before']} -> {item['after']}"
        for item in changes
    ]
    lines.extend(
        [
            f"[KW {item['index']}] pulled_from={item['pulled_from']} | {item['before']} -> {item['after']}"
            for item in keyword_changes
        ]
    )

    (logs_dir / f"{target.stem}.annotation.log").write_text(
        "\n".join(lines),
        encoding="utf-8",
    )

    return target


def main() -> None:
    parser = argparse.ArgumentParser(description="Normalize annotation markers in draft")
    parser.add_argument("draft", type=Path, help="Draft docx")
    parser.add_argument("--output", type=Path, default=None, help="Output docx")
    parser.add_argument("--logs-dir", type=Path, default=None, help="Where to write logs/reports")
    args = parser.parse_args()

    if not args.draft.exists():
        raise SystemExit("Draft not found")

    out = normalize_annotations(
        args.draft,
        output_path=args.output,
        logs_dir=args.logs_dir,
    )

    print(f"OK: {out}")


if __name__ == "__main__":
    main()
