from __future__ import annotations

import re
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Iterable

from docx import Document

UDC_RE = re.compile(r"^\s*(УДК|UDC)\b", re.IGNORECASE)
BODY_MARKER_RE = re.compile(r"^\s*(Анотація|Abstract|Ключові слова|Keywords|Вступ|Introduction)\b", re.IGNORECASE)

AFFILIATION_KEYWORDS = (
    "університет",
    "інститут",
    "академ",
    "кафедра",
    "факультет",
    "націон",
    "держав",
    "центр",
    "лаборатор",
    "відділ",
    "university",
    "institute",
    "academy",
    "faculty",
    "department",
    "college",
    "research",
    "center",
    "centre",
    "laboratory",
)


@dataclass(slots=True)
class LineDecision:
    text: str
    label: str
    method: str
    score: float
    bold_ratio: float
    font_size: float | None
    position: int
    style_name: str


def _normalize(text: str) -> str:
    value = re.sub(r"\s+", " ", (text or "").casefold()).strip()
    value = re.sub(r"[^\w\s]", " ", value, flags=re.UNICODE)
    return re.sub(r"\s+", " ", value).strip()


def _similar(a: str, b: str) -> float:
    return SequenceMatcher(None, _normalize(a), _normalize(b)).ratio()


def _looks_like_affiliation(text: str) -> bool:
    low = text.casefold()
    return any(key in low for key in AFFILIATION_KEYWORDS)


def _looks_like_author_line(text: str) -> bool:
    if len(text) > 120 or BODY_MARKER_RE.search(text) or UDC_RE.search(text):
        return False
    lowered = text.casefold()
    if any(marker in lowered for marker in ("orcid", "doi", "@")):
        return False
    if any(char.isdigit() for char in text):
        return False
    words = [word for word in re.split(r"\s+", text) if word]
    if len(words) < 2 or len(words) > 5:
        return False
    titlecase_words = sum(1 for word in words if word[:1].isupper())
    return titlecase_words >= 2


def _extract_probable_authors(header_lines: list[str]) -> list[str]:
    candidates: list[str] = []
    seen: set[str] = set()
    for line in header_lines:
        parts = [part for part in re.split(r"[,;]+", line) if part.strip()]
        ordered_parts = parts if parts else [line]
        for part in ordered_parts:
            candidate = re.sub(r"\s+", " ", part).strip().strip(",.;: ")
            if not _looks_like_author_line(candidate):
                continue
            key = candidate.casefold()
            if key in seen:
                continue
            seen.add(key)
            candidates.append(candidate)
    return candidates


def _line_features(doc: Document, header_lines: list[str]) -> dict[str, tuple[float, float | None, str]]:
    features: dict[str, tuple[float, float | None, str]] = {}
    header_set = {_normalize(line) for line in header_lines if line.strip()}
    for para in doc.paragraphs[:60]:
        text = (para.text or "").strip()
        if not text:
            continue
        key = _normalize(text)
        if key not in header_set:
            continue
        bold_runs = 0
        total_runs = 0
        sizes: list[float] = []
        for run in para.runs:
            total_runs += 1
            if run.bold:
                bold_runs += 1
            if run.font.size is not None:
                try:
                    sizes.append(float(run.font.size.pt))
                except Exception:
                    pass
        bold_ratio = (bold_runs / total_runs) if total_runs else 0.0
        font_size = max(sizes) if sizes else None
        style_name = getattr(para.style, "name", "") or ""
        features[key] = (bold_ratio, font_size, style_name)
    return features


def _style_bonus(style_name: str, label: str) -> float:
    if not style_name:
        return 0.0
    key = style_name.casefold()
    if label == "person" and ("heading 2" in key or "заголовок 2" in key or "autor" in key):
        return 0.12
    if label == "title" and ("heading 3" in key or "заголовок 3" in key or "назва1" in key):
        return 0.08
    return 0.0


def _decide_lines(header_lines: list[str], probable_authors: set[str], features: dict[str, tuple[float, float | None, str]]) -> list[LineDecision]:
    decisions: list[LineDecision] = []
    for idx, line in enumerate(header_lines):
        text = line.strip()
        if not text:
            continue
        key = _normalize(text)
        bold_ratio, font_size, style_name = features.get(key, (0.0, None, ""))
        format_bonus = 0.0
        if bold_ratio >= 0.5:
            format_bonus += 0.1
        if font_size and font_size >= 14:
            format_bonus += 0.05

        if text in probable_authors:
            score = 0.8 + format_bonus + _style_bonus(style_name, "person")
            decisions.append(LineDecision(text=text, label="person", method="probable", score=score, bold_ratio=bold_ratio, font_size=font_size, position=idx, style_name=style_name))
            continue
        if _looks_like_author_line(text):
            score = 0.7 + format_bonus + _style_bonus(style_name, "person")
            decisions.append(LineDecision(text=text, label="person", method="pattern", score=score, bold_ratio=bold_ratio, font_size=font_size, position=idx, style_name=style_name))
            continue
        if text.isupper() and 10 <= len(text) <= 140:
            score = 0.65 + format_bonus + _style_bonus(style_name, "title")
            decisions.append(LineDecision(text=text, label="title", method="caps", score=score, bold_ratio=bold_ratio, font_size=font_size, position=idx, style_name=style_name))
            continue
        if _looks_like_affiliation(text):
            score = 0.6 + format_bonus
            decisions.append(LineDecision(text=text, label="org", method="affiliation", score=score, bold_ratio=bold_ratio, font_size=font_size, position=idx, style_name=style_name))
            continue
        decisions.append(LineDecision(text=text, label="other", method="fallback", score=0.2, bold_ratio=bold_ratio, font_size=font_size, position=idx, style_name=style_name))
    return decisions


def _collect_entities(decisions: Iterable[LineDecision]) -> tuple[list[str], list[str]]:
    people: list[str] = []
    orgs: list[str] = []
    seen_people: set[str] = set()
    seen_orgs: set[str] = set()
    for item in decisions:
        if item.label == "person":
            key = item.text.casefold()
            if key not in seen_people:
                seen_people.add(key)
                people.append(item.text)
        if item.label == "org":
            key = item.text.casefold()
            if key not in seen_orgs:
                seen_orgs.add(key)
                orgs.append(item.text)
    return people, orgs


def detect_header(path: Path, *, max_lines: int = 20) -> dict[str, object]:
    try:
        doc = Document(path)
    except Exception:
        return {"header_lines": [], "authors": [], "orgs": [], "decisions": []}
    header_lines: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        if BODY_MARKER_RE.search(text):
            break
        header_lines.append(text)
        if len(header_lines) >= max_lines:
            break

    probable = set(_extract_probable_authors(header_lines))
    features = _line_features(doc, header_lines)
    decisions = _decide_lines(header_lines, probable, features)
    authors, orgs = _collect_entities(decisions)

    return {
        "header_lines": header_lines,
        "authors": authors,
        "orgs": orgs,
        "decisions": [
            {
                "text": item.text,
                "label": item.label,
                "method": item.method,
                "score": item.score,
                "position": item.position,
            }
            for item in decisions
        ],
    }
