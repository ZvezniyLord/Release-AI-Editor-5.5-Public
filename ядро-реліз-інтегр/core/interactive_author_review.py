from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document

if __package__ in {None, ""}:
    ROOT = Path(__file__).resolve().parents[1]
    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))

try:
    from core.normalize_authors import _collect_author_candidates
except Exception:
    from normalize_authors import _collect_author_candidates  # type: ignore


def _context_lines(doc: Document, idx: int, span: int = 2) -> tuple[str, str]:
    before = ""
    after = ""
    # idx is 1-based
    i = idx - 2
    steps = 0
    while i >= 0 and steps < span:
        t = (doc.paragraphs[i].text or "").strip()
        if t:
            before = t
            steps += 1
        i -= 1
    i = idx
    steps = 0
    while i < len(doc.paragraphs) and steps < span:
        t = (doc.paragraphs[i].text or "").strip()
        if t:
            after = t
            steps += 1
        i += 1
    return before, after


def _review_doc(doc_path: Path, output_path: Path, *, in_place: bool) -> dict:
    doc = Document(doc_path)
    candidates = _collect_author_candidates(doc)
    candidates_total = len(candidates)
    accepted = 0
    rejected = 0
    skipped = 0
    applied = 0
    accept_all = False
    stop_file = False

    for idx in candidates:
        if stop_file:
            break
        if idx - 1 >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx - 1]
        text = (para.text or "").strip()
        if not text:
            continue
        if accept_all:
            try:
                para.style = "AUTOR"
                applied += 1
                accepted += 1
            except Exception:
                skipped += 1
            continue

        before, after = _context_lines(doc, idx)
        print("\n---")
        print(f"Файл: {doc_path.name}")
        print(f"Абзац #{idx}")
        if before:
            print(f"До: {before}")
        print(f"Кандидат: {text}")
        if after:
            print(f"Після: {after}")
        print("Дії: [y] так, [n] ні, [a] так для всіх у цьому файлі, [s] пропустити файл, [q] вихід")

        while True:
            ans = input("> ").strip().lower()
            if ans == "y":
                try:
                    para.style = "AUTOR"
                    applied += 1
                    accepted += 1
                except Exception:
                    skipped += 1
                break
            if ans == "n":
                rejected += 1
                break
            if ans == "a":
                accept_all = True
                try:
                    para.style = "AUTOR"
                    applied += 1
                    accepted += 1
                except Exception:
                    skipped += 1
                break
            if ans == "s":
                remaining = len([c for c in candidates if c >= idx])
                skipped += remaining
                stop_file = True
                break
            if ans == "q":
                raise KeyboardInterrupt
            print("Невірна команда. Введи y/n/a/s/q.")

    if in_place:
        doc.save(doc_path)
        out_path = doc_path
    else:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        out_path = output_path

    return {
        "file": str(doc_path),
        "output": str(out_path),
        "candidates": candidates_total,
        "accepted": accepted,
        "rejected": rejected,
        "skipped": skipped,
        "applied": applied,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Interactive author review on normalized articles")
    parser.add_argument("input_folder", type=Path, help="Папка з нормалізованими статтями (Статті_норм)")
    parser.add_argument("--output-folder", type=Path, default=None, help="Куди зберігати (за замовчуванням *_authors)")
    parser.add_argument("--in-place", action="store_true", help="Перезаписувати файли на місці")
    parser.add_argument("--report", type=Path, default=None, help="JSON-звіт")
    args = parser.parse_args()

    source = args.input_folder
    if not source.exists():
        raise SystemExit(f"Папку не знайдено: {source}")

    if args.in_place:
        output_folder = source
    else:
        output_folder = args.output_folder or source.parent / f"{source.name}_authors"

    report_path = args.report or (output_folder / "interactive_authors_report.json")

    items = []
    try:
        files = sorted(source.glob("*.docx"))
        for doc_path in files:
            out_path = output_folder / doc_path.name
            item = _review_doc(doc_path, out_path, in_place=args.in_place)
            items.append(item)
    except KeyboardInterrupt:
        print("Зупинено користувачем.")

    report = {
        "input_folder": str(source),
        "output_folder": str(output_folder),
        "items": items,
    }
    output_folder.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Звіт: {report_path}")


if __name__ == "__main__":
    main()
