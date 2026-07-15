from __future__ import annotations

import argparse
import os
import sys
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

if __package__ in {None, ""}:
    ROOT = Path(__file__).resolve().parents[1]
    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))

MAX_HEADER_PARAS = 35
MAX_LINE_LENGTH = 180
MIN_LINE_LEN = 3
MODULE_ROOT = Path(__file__).resolve().parents[1]
MARKERS_PATH = MODULE_ROOT / "resources" / "markers.json"


def _load_markers() -> dict[str, list[str]]:
    default = {
        "author_stop_markers": [],
        "author_exclude_markers": [],
    }
    try:
        data = json.loads(MARKERS_PATH.read_text(encoding="utf-8"))
    except Exception:
        return default
    if not isinstance(data, dict):
        return default
    def _norm_list(key: str) -> list[str]:
        raw = data.get(key, [])
        if not isinstance(raw, list):
            return []
        out: list[str] = []
        for item in raw:
            value = str(item).casefold().strip()
            if value:
                out.append(value)
        return out

    stop_sources = (
        _norm_list("annotation_markers")
        + _norm_list("keyword_markers")
        + _norm_list("reference_markers")
        + _norm_list("author_stop_markers")
    )
    # Stable de-duplication while preserving first appearance order.
    seen: set[str] = set()
    stop_markers: list[str] = []
    for marker in stop_sources:
        if marker in seen:
            continue
        seen.add(marker)
        stop_markers.append(marker)

    return {
        "author_stop_markers": stop_markers,
        "author_exclude_markers": _norm_list("author_exclude_markers"),
    }


def _looks_like_author_line(text: str) -> bool:
    if len(text) > 120:
        return False
    if any(char.isdigit() for char in text):
        return False
    lowered = text.casefold()
    if any(marker in lowered for marker in _load_markers()["author_exclude_markers"]):
        return False
    words = [word for word in text.split() if word]
    if len(words) < 2 or len(words) > 5:
        return False
    def _is_titlecase_word(word: str) -> bool:
        token = word.strip("*").strip(",.;:()[]{}")
        return bool(token) and token[:1].isupper()

    titlecase_words = sum(1 for word in words if _is_titlecase_word(word))
    return titlecase_words >= 2


def _is_author_stop_marker(text: str) -> bool:
    raw = (text or "").strip().casefold()
    if not raw:
        return False
    raw = re.sub(r"[\s:;\-–—]+", " ", raw)
    raw = re.sub(r"\s+", " ", raw).strip()
    triggers = _load_markers()["author_stop_markers"]
    return any(raw.startswith(t) for t in triggers)


def _expected_author_count(draft_path: Path) -> int:
    matches_path = draft_path.parent / "matches.json"
    if not matches_path.exists():
        return 0
    try:
        data = json.loads(matches_path.read_text(encoding="utf-8"))
    except Exception:
        return 0
    if not isinstance(data, list):
        return 0
    count = 0
    for item in data:
        if not isinstance(item, dict):
            continue
        method = (item.get("match_method") or "").strip()
        if method in {"missing", "missing_source_file", "free_listener"}:
            continue
        authors = item.get("authors") or []
        if isinstance(authors, list) and authors:
            count += len(authors)
        else:
            count += 1
    return count


def _collect_author_candidates(doc: Document) -> list[int]:
    candidates: list[int] = []
    since_break = 0
    allow_author = True

    def has_page_break(p) -> bool:
        for run in p.runs:
            el = run._element
            if el is None:
                continue
            for node in el.iter():
                if node.tag.endswith("}br") and node.get(qn("w:type")) == "page":
                    return True
                if node.tag.endswith("}lastRenderedPageBreak"):
                    return True
        return False

    for idx, para in enumerate(doc.paragraphs, start=1):
        raw_text = para.text or ""
        if "\x0c" in raw_text or has_page_break(para):
            since_break = 0
            allow_author = True
            continue
        since_break += 1
        text = raw_text.strip()
        if not text:
            continue
        if _is_author_stop_marker(text):
            allow_author = False
            continue
        if len(text) < MIN_LINE_LEN or len(text) > MAX_LINE_LENGTH:
            continue
        if since_break > MAX_HEADER_PARAS:
            continue
        is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        if allow_author and _looks_like_author_line(text) and (is_centered or text.isupper() or text.istitle()):
            candidates.append(idx)
    return candidates


def normalize_authors_docx(
    draft_path: Path,
    *,
    output_path: Path | None = None,
) -> Path:
    doc = Document(draft_path)
    candidates = _collect_author_candidates(doc)

    report = {
        "authors_expected": _expected_author_count(draft_path),
        "authors_detected": len(candidates),
    }
    try:
        (draft_path.parent / "author_detection_report.json").write_text(
            json.dumps(report, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception:
        pass

    for idx in candidates:
        if idx - 1 >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx - 1]
        try:
            para.style = "AUTOR"
        except Exception:
            pass

    target = output_path or draft_path
    doc.save(target)
    return target


def main() -> None:
    parser = argparse.ArgumentParser(description="Normalize authors in draft docx")
    parser.add_argument("draft", type=Path, help="Draft docx")
    parser.add_argument("--output", type=Path, default=None, help="Output docx (copy)")
    args = parser.parse_args()

    if not args.draft.exists():
        raise SystemExit("Draft not found")
    out = normalize_authors_docx(args.draft, output_path=args.output)
    try:
        os.startfile(str(out))
    except Exception:
        pass


if __name__ == "__main__":
    main()
