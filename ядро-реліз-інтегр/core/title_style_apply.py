from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from core.align_titles_adapter import _maybe_align_titles
from core.title_detection import (
    _collect_title_candidates,
    _load_matches_by_title,
    _load_titles_from_matches,
    _merge_title_paragraphs,
    _normalize_key,
    _normalize_key_compact,
)

try:
    from core.paragraph_utils import ensure_single_empty_around
except Exception:
    from paragraph_utils import ensure_single_empty_around


def _strip_outer_title_quotes(text: str) -> str:
    value = (text or "").strip()
    if len(value) < 2:
        return value
    quote_pairs = [
        ('"', '"'),
        ("“", "”"),
        ("«", "»"),
    ]
    for left, right in quote_pairs:
        if value.startswith(left) and value.endswith(right):
            inner = value[1:-1].strip()
            return inner or value
    return value


def normalize_titles_docx(
    draft_path: Path,
    *,
    output_path: Path | None = None,
    align_titles: bool = False,
    align_threshold: float = 0.9,
    align_auto: bool = False,
    logs_dir: Path | None = None,
    titles_override: list[str] | None = None,
    use_heuristics: bool = True,
    run_dir: Path | None = None,
    write_report: bool = True,
) -> Path:
    align_report = None
    if align_titles:
        align_report = _maybe_align_titles(
            draft_path,
            align_threshold,
            align_auto,
            run_dir=run_dir,
            titles_override=titles_override,
        )
    doc = Document(draft_path)

    if titles_override is not None:
        titles = [t for t in titles_override if t and str(t).strip()]
        # Якщо інтерактив змінив назву — беремо оновлену.
        if align_report and align_report.get("changes"):
            for change in align_report["changes"]:
                if change.get("from") in titles:
                    titles = [change.get("to") or change.get("from")]
                    break
        match_stats = {"total": len(titles), "title_match": len(titles), "free_listener": 0, "missing": 0, "other": 0}
        if titles_override and use_heuristics:
            use_heuristics = False
    else:
        titles, match_stats = _load_titles_from_matches(draft_path)
    title_index = {_normalize_key(t): t for t in titles if _normalize_key(t)}
    title_index_compact = {_normalize_key_compact(t): t for t in titles if _normalize_key_compact(t)}
    title_keys = set(title_index.keys())
    title_keys_compact = set(title_index_compact.keys())
    title_hits: set[int] = set()
    merged_counts: dict[str, int] = {"2": 0, "3": 0, "4": 0, "5": 0, "6": 0}

    if title_keys or title_keys_compact:
        merged_hits, merged_counts = _merge_title_paragraphs(doc, title_index, title_keys_compact)
        title_hits |= merged_hits
        for idx, para in enumerate(doc.paragraphs, start=1):
            text = para.text or ""
            key = _normalize_key(text)
            keyc = _normalize_key_compact(text)
            if (key and key in title_keys) or (keyc and keyc in title_keys_compact):
                title_hits.add(idx)

    candidates = _collect_title_candidates(doc) if use_heuristics else {"title": []}

    title_hits_keys = {_normalize_key(doc.paragraphs[i - 1].text) for i in title_hits}
    title_hits_keys_compact = {_normalize_key_compact(doc.paragraphs[i - 1].text) for i in title_hits}
    missing_titles = [
        t
        for t in titles
        if _normalize_key(t) not in title_hits_keys and _normalize_key_compact(t) not in title_hits_keys_compact
    ]
    missing_details = []
    if missing_titles:
        matches_by_title = _load_matches_by_title(draft_path)
        for title in missing_titles:
            item = matches_by_title.get(title, {})
            authors = item.get("authors") if isinstance(item, dict) else None
            source_path = ""
            source_label = ""
            if isinstance(item, dict):
                source_path = str(
                    item.get("cleaned_path")
                    or item.get("relocated_path")
                    or item.get("matched_path")
                    or ""
                )
                if source_path:
                    source_label = Path(source_path).stem
            missing_details.append(
                {
                    "title": title,
                    "authors": authors or [],
                    "source_label": source_label,
                    "source_path": source_path,
                }
            )
    report = {
        "titles_in_json": len(titles),
        "titles_matched": len(title_hits),
        "titles_missing": missing_titles,
        "titles_missing_details": missing_details,
        "merged_counts": merged_counts,
        "match_method_stats": match_stats,
    }
    if write_report:
        try:
            if logs_dir is None:
                logs_dir = draft_path.parent
            logs_dir.mkdir(parents=True, exist_ok=True)
            (logs_dir / "title_detection_report.json").write_text(
                json.dumps(report, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        except Exception:
            pass

    for idx in sorted(title_hits):
        if idx - 1 >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx - 1]
        cleaned = _strip_outer_title_quotes(para.text or "")
        if cleaned != (para.text or ""):
            para.text = cleaned
        try:
            para.style = "Назва1"
        except Exception:
            pass
        ensure_single_empty_around(para)

    for idx in candidates["title"]:
        if idx in title_hits:
            continue
        if idx - 1 >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx - 1]
        cleaned = _strip_outer_title_quotes(para.text or "")
        if cleaned != (para.text or ""):
            para.text = cleaned
        try:
            para.style = "Назва1"
        except Exception:
            pass
        ensure_single_empty_around(para)

    target = output_path or draft_path
    doc.save(target)
    return target
