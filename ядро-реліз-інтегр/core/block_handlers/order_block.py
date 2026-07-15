from __future__ import annotations

import json
import re
from pathlib import Path
from difflib import SequenceMatcher

from docx import Document

UDC_RE = re.compile(r"^\s*(\u0423\u0414\u041a|UDC|UDK)\b", re.IGNORECASE)
DOI_RE = re.compile(r"^\s*(doi\s*[:\s]|https?\s*:\s*//\s*doi\.org/)", re.IGNORECASE)
BODY_START_RE = re.compile(
    r"^\s*(\u0432\u0441\u0442\u0443\u043f|\u043c\u0435\u0442\u0430|\u043e\u0441\u043d\u043e\u0432\u043d\u0430 \u0447\u0430\u0441\u0442\u0438\u043d\u0430|\u043c\u0430\u0442\u0435\u0440\u0456\u0430\u043b\u0438|\u0440\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442\u0438|\u043e\u0431\u0433\u043e\u0432\u043e\u0440\u0435\u043d\u043d\u044f|\u0432\u0438\u0441\u043d\u043e\u0432|introduction|aim|results|conclusions)\b",
    re.IGNORECASE,
)
ANNOTATION_RE = re.compile(r"^\s*(\u0430\u043d\u043e\u0442\u0430\u0446\u0456\u044f|annotation|abstract)\b", re.IGNORECASE)
KEYWORDS_RE = re.compile(r"^\s*(\u043a\u043b\u044e\u0447\u043e\u0432\u0456 \u0441\u043b\u043e\u0432\u0430|keywords|key words)\b", re.IGNORECASE)
REFERENCES_RE = re.compile(
    r"^\s*(\u0441\u043f\u0438\u0441\u043e\u043a \u0432\u0438\u043a\u043e\u0440\u0438\u0441\u0442\u0430\u043d\u0438\u0445 \u0434\u0436\u0435\u0440\u0435\u043b|references|bibliography|list of references|list of reference)\b",
    re.IGNORECASE,
)
FRONT_SCOPE_MAX_PARAGRAPHS = 45


def _norm(s: str) -> str:
    s = (s or "").replace("\u00a0", " ").strip()
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"[\.\,\;\:\!\?]+$", "", s)
    return s.casefold()


def _front_scope_end_idx(doc: Document) -> int:
    total = len(doc.paragraphs)
    if total <= 0:
        return 0
    return min(total, FRONT_SCOPE_MAX_PARAGRAPHS)


def _find_title_idx(doc: Document, expected_title: str, *, max_idx: int | None = None) -> int | None:
    limit = max_idx if isinstance(max_idx, int) and max_idx > 0 else len(doc.paragraphs)

    def _is_disallowed_title_text(text: str) -> bool:
        t = (text or "").strip()
        return bool(ANNOTATION_RE.search(t) or KEYWORDS_RE.search(t) or REFERENCES_RE.search(t) or BODY_START_RE.search(t))

    style_hits: list[int] = []
    for idx, para in enumerate(doc.paragraphs[:limit], start=1):
        para_text = (para.text or "").strip()
        if not para_text or _is_disallowed_title_text(para_text):
            continue
        style_name = ""
        try:
            style_name = (para.style.name or "").casefold()
        except Exception:
            style_name = ""
        if any(token in style_name for token in ("назва1", "title", "heading 3")):
            style_hits.append(idx)
    if style_hits:
        return style_hits[0]

    target = _norm(expected_title)
    if not target:
        return None
    best_idx: int | None = None
    best_score = 0.0
    for idx, para in enumerate(doc.paragraphs[:limit], start=1):
        para_text = (para.text or "").strip()
        if not para_text or _is_disallowed_title_text(para_text):
            continue
        text_norm = _norm(para_text)
        if text_norm == target:
            return idx
        score = SequenceMatcher(None, text_norm, target).ratio()
        if score > best_score:
            best_score = score
            best_idx = idx
    if best_idx is not None and best_score >= 0.72:
        return best_idx
    return None


def _find_title_fallback_idx(doc: Document, *, max_idx: int | None = None) -> int | None:
    limit = max_idx if isinstance(max_idx, int) and max_idx > 0 else len(doc.paragraphs)
    for idx, para in enumerate(doc.paragraphs[:limit], start=1):
        text = (para.text or "").strip()
        if not text:
            continue
        if _is_udc_like(text):
            continue
        if ANNOTATION_RE.search(text) or KEYWORDS_RE.search(text) or REFERENCES_RE.search(text) or BODY_START_RE.search(text):
            continue

        style_name = ""
        try:
            style_name = (para.style.name or "").casefold()
        except Exception:
            style_name = ""
        if any(token in style_name for token in ("назва1", "title", "heading 3")):
            return idx

        letters = [ch for ch in text if ch.isalpha()]
        if len(letters) < 18 or len(text) > 220:
            continue
        upper_ratio = sum(1 for ch in letters if ch.isupper()) / len(letters)
        if upper_ratio >= 0.65:
            return idx
    return None


def _is_udc_like(text: str) -> bool:
    value = (text or "").strip()
    return bool(UDC_RE.search(value) or DOI_RE.search(value))


def _find_udc_idx(doc: Document, *, max_idx: int | None = None) -> int | None:
    limit = max_idx if isinstance(max_idx, int) and max_idx > 0 else len(doc.paragraphs)
    for idx, para in enumerate(doc.paragraphs[:limit], start=1):
        text = (para.text or "").strip()
        if _is_udc_like(text):
            return idx
    return None


def _collect_udc_indices(doc: Document, first_udc_idx: int, *, max_idx: int | None = None) -> list[int]:
    limit = max_idx if isinstance(max_idx, int) and max_idx > 0 else len(doc.paragraphs)
    if first_udc_idx < 1 or first_udc_idx > limit:
        return []
    indices: list[int] = [first_udc_idx]
    idx = first_udc_idx + 1
    while idx <= min(limit, len(doc.paragraphs)):
        text = (doc.paragraphs[idx - 1].text or "").strip()
        if not text:
            idx += 1
            continue
        if _is_udc_like(text):
            indices.append(idx)
            idx += 1
            continue
        break
    return indices


def _collect_header_after_title(doc: Document, title_idx: int, *, max_idx: int | None = None) -> list[int]:
    limit = max_idx if isinstance(max_idx, int) and max_idx > 0 else len(doc.paragraphs)
    header: list[int] = []
    for idx in range(title_idx + 1, min(limit, len(doc.paragraphs)) + 1):
        text = (doc.paragraphs[idx - 1].text or "").strip()
        if not text:
            if header:
                break
            continue
        if BODY_START_RE.search(text):
            break
        if ANNOTATION_RE.search(text) or KEYWORDS_RE.search(text) or REFERENCES_RE.search(text):
            break
        header.append(idx)
    return header


def _find_front_boundary_idx(doc: Document, *, max_idx: int | None = None) -> int | None:
    limit = max_idx if isinstance(max_idx, int) and max_idx > 0 else len(doc.paragraphs)
    for idx, para in enumerate(doc.paragraphs[:limit], start=1):
        text = (para.text or "").strip()
        if not text:
            continue
        if ANNOTATION_RE.search(text) or KEYWORDS_RE.search(text) or BODY_START_RE.search(text) or REFERENCES_RE.search(text):
            return idx
    return None


def _find_first_non_header_idx(
    doc: Document,
    scan_start_idx: int,
    *,
    stop_before_idx: int | None = None,
    max_idx: int | None = None,
) -> int | None:
    limit = max_idx if isinstance(max_idx, int) and max_idx > 0 else len(doc.paragraphs)
    upper = min(limit, len(doc.paragraphs))
    if stop_before_idx is not None:
        upper = min(upper, stop_before_idx - 1)
    for idx in range(max(1, scan_start_idx), upper + 1):
        text = (doc.paragraphs[idx - 1].text or "").strip()
        if not text:
            continue
        if _is_udc_like(text):
            continue
        if ANNOTATION_RE.search(text) or KEYWORDS_RE.search(text) or REFERENCES_RE.search(text) or BODY_START_RE.search(text):
            return idx
        if not _is_header_candidate(text):
            return idx
    return None


def _effective_front_boundary_idx(
    doc: Document,
    scan_start_idx: int,
    title_idx: int | None,
    *,
    max_idx: int | None = None,
) -> int | None:
    marker_boundary = _find_front_boundary_idx(doc, max_idx=max_idx)
    prose_boundary = _find_first_non_header_idx(
        doc,
        scan_start_idx,
        stop_before_idx=title_idx,
        max_idx=max_idx,
    )
    candidates = [value for value in (marker_boundary, prose_boundary) if value is not None]
    if title_idx is not None:
        candidates.append(title_idx)
    return min(candidates) if candidates else None


def _is_header_candidate(text: str) -> bool:
    value = (text or "").strip()
    if not value:
        return False
    if _is_udc_like(value):
        return False
    if ANNOTATION_RE.search(value) or KEYWORDS_RE.search(value) or REFERENCES_RE.search(value) or BODY_START_RE.search(value):
        return False
    # Front-matter header lines are typically short and compact. Long prose
    # paragraphs (annotation/body) must not be treated as header rows.
    if len(value) > 180:
        return False
    return True


def _collect_header_indices(doc: Document, scan_start_idx: int, title_idx: int, *, max_idx: int | None = None) -> list[int]:
    limit = max_idx if isinstance(max_idx, int) and max_idx > 0 else len(doc.paragraphs)
    boundary = _effective_front_boundary_idx(doc, scan_start_idx, title_idx, max_idx=limit) or (min(limit, len(doc.paragraphs)) + 1)
    scan_from = max(scan_start_idx, 1)
    scan_to = min(boundary - 1, min(limit, len(doc.paragraphs)))
    if scan_from > scan_to:
        return []
    compact: list[int] = []
    for idx in range(scan_from, scan_to + 1):
        if idx == title_idx:
            continue
        para = doc.paragraphs[idx - 1]
        text = (para.text or "").strip()
        if not text:
            continue
        if not _is_header_candidate(text):
            continue
        compact.append(idx)
    return compact


def _ensure_blank_after_top_udc_block(doc: Document) -> bool:
    """Force one blank line after the first UDC/DOI block in front matter."""
    first_udc = _find_udc_idx(doc, max_idx=_front_scope_end_idx(doc))
    if first_udc is None:
        return False
    udc_block = _collect_udc_indices(doc, first_udc, max_idx=_front_scope_end_idx(doc)) or [first_udc]
    return _ensure_single_blank_after(doc, udc_block[-1])


def _apply_udc_style_to_block(doc: Document, indices: list[int]) -> bool:
    changed = False
    for idx in indices:
        if idx < 1 or idx > len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx - 1]
        try:
            style_name = para.style.name or ""
        except Exception:
            style_name = ""
        if style_name == "UDC":
            continue
        try:
            para.style = "UDC"
            changed = True
        except Exception:
            pass
    return changed


def _remove_paragraph(para) -> None:
    p = para._element
    p.getparent().remove(p)


def _remove_empty_paragraphs_in_range(doc: Document, start_idx: int, end_idx: int) -> bool:
    if start_idx > end_idx:
        return False
    changed = False
    upper = min(end_idx, len(doc.paragraphs))
    lower = max(1, start_idx)
    for idx in range(upper, lower - 1, -1):
        para = doc.paragraphs[idx - 1]
        if (para.text or "").strip():
            continue
        _remove_paragraph(para)
        changed = True
    return changed


def _ensure_single_blank_after(doc: Document, idx: int, *, style_name: str | None = None) -> bool:
    if idx < 1 or idx > len(doc.paragraphs):
        return False
    changed = False

    if idx >= len(doc.paragraphs):
        blank = doc.add_paragraph("")
        if style_name:
            try:
                blank.style = style_name
            except Exception:
                pass
        return True

    next_para = doc.paragraphs[idx]
    if (next_para.text or "").strip():
        blank = next_para.insert_paragraph_before("")
        if style_name:
            try:
                blank.style = style_name
            except Exception:
                pass
        changed = True
    else:
        if style_name:
            try:
                next_para.style = style_name
            except Exception:
                pass

    while idx + 1 < len(doc.paragraphs):
        after = doc.paragraphs[idx + 1]
        if (after.text or "").strip():
            break
        _remove_paragraph(after)
        changed = True

    return changed


def _recompute_front_indices(doc: Document, expected_title: str) -> tuple[list[int], int | None, list[int]]:
    scope_end = _front_scope_end_idx(doc)
    udc_idx = _find_udc_idx(doc, max_idx=scope_end)
    udc_indices: list[int] = []
    if udc_idx is not None:
        udc_indices = _collect_udc_indices(doc, udc_idx, max_idx=scope_end) or [udc_idx]

    title_idx = _find_title_idx(doc, expected_title, max_idx=scope_end)
    header_indices: list[int] = []
    if title_idx is not None and udc_indices:
        header_indices = [
            idx
            for idx in range(max(udc_indices) + 1, title_idx)
            if (doc.paragraphs[idx - 1].text or "").strip()
        ]
    return udc_indices, title_idx, header_indices


def enforce_primary_order_with_title(doc_path: Path, expected_title: str, logs_dir: Path | None = None) -> dict:
    doc = Document(doc_path)
    scope_end_idx = _front_scope_end_idx(doc)
    title_idx = _find_title_idx(doc, expected_title, max_idx=scope_end_idx)
    title_reason = "expected_title"
    if title_idx is None:
        title_idx = _find_title_fallback_idx(doc, max_idx=scope_end_idx)
        if title_idx is not None:
            title_reason = "fallback_detected"

    if title_idx is None:
        # Even when title matching fails, keep basic front-matter spacing sane.
        changed = _ensure_blank_after_top_udc_block(doc)
        if changed:
            doc.save(doc_path)
        result = {
            "changed": changed,
            "reason": "title_not_found",
            "title": expected_title,
            "scope": "first_n_paragraphs",
            "scope_max_paragraphs": FRONT_SCOPE_MAX_PARAGRAPHS,
            "scope_end_idx": scope_end_idx,
        }
        if logs_dir is not None:
            logs_dir.mkdir(parents=True, exist_ok=True)
            (logs_dir / f"{doc_path.stem}.order_report.json").write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        return result

    udc_idx = _find_udc_idx(doc, max_idx=scope_end_idx)
    if udc_idx is None:
        p = doc.add_paragraph("\u0423\u0414\u041a")
        try:
            p.style = "UDC"
        except Exception:
            pass
        udc_idx = len(doc.paragraphs)
    udc_indices = _collect_udc_indices(doc, udc_idx, max_idx=scope_end_idx)
    if not udc_indices:
        udc_indices = [udc_idx]

    body = doc._element.body
    header_indices = _collect_header_indices(doc, max(udc_indices) + 1, title_idx, max_idx=scope_end_idx)
    if not header_indices:
        header_indices = _collect_header_after_title(doc, title_idx, max_idx=scope_end_idx)

    moved_elements = [doc.paragraphs[i - 1]._element for i in udc_indices if 1 <= i <= len(doc.paragraphs)]
    moved_elements.extend(doc.paragraphs[i - 1]._element for i in header_indices if 1 <= i <= len(doc.paragraphs))
    if 1 <= title_idx <= len(doc.paragraphs):
        moved_elements.append(doc.paragraphs[title_idx - 1]._element)

    dedup = []
    seen = set()
    for elem in moved_elements:
        key = id(elem)
        if key in seen:
            continue
        seen.add(key)
        dedup.append(elem)
    moved_elements = dedup

    top = [p._element for p in doc.paragraphs[: len(moved_elements)]]
    changed = top != moved_elements

    if changed:
        for elem in moved_elements:
            try:
                body.remove(elem)
            except Exception:
                pass
        for elem in reversed(moved_elements):
            body.insert(0, elem)

    front_boundary_after_move = _find_front_boundary_idx(doc, max_idx=_front_scope_end_idx(doc))
    if front_boundary_after_move and front_boundary_after_move > 1:
        changed = _remove_empty_paragraphs_in_range(doc, 1, front_boundary_after_move - 1) or changed

    # Normalize front-matter spacing:
    # UDC/DOI block -> blank -> authors/affiliation -> blank -> title -> blank.
    # Recompute key indices after potential move to avoid unstable XML object identity.
    scope_end_after = _front_scope_end_idx(doc)
    udc_idx_after = _find_udc_idx(doc, max_idx=scope_end_after)
    udc_indices_after: list[int] = []
    if udc_idx_after is not None:
        udc_indices_after = _collect_udc_indices(doc, udc_idx_after, max_idx=scope_end_after)
        if not udc_indices_after:
            udc_indices_after = [udc_idx_after]

    title_idx_after = _find_title_idx(doc, expected_title, max_idx=scope_end_after)
    header_indices_after: list[int] = []
    if title_idx_after is not None and udc_indices_after:
        header_indices_after = _collect_header_indices(
            doc,
            max(udc_indices_after) + 1,
            title_idx_after,
            max_idx=scope_end_after,
        )
        if not header_indices_after:
            header_indices_after = _collect_header_after_title(doc, title_idx_after, max_idx=scope_end_after)

    if title_idx_after is not None and udc_indices_after:
        changed = _remove_empty_paragraphs_in_range(doc, udc_indices_after[-1] + 1, title_idx_after - 1) or changed
        scope_end_after = _front_scope_end_idx(doc)
        title_idx_after = _find_title_idx(doc, expected_title, max_idx=scope_end_after)
        if title_idx_after is not None:
            front_boundary = _effective_front_boundary_idx(
                doc,
                max(udc_indices_after) + 1,
                title_idx_after,
                max_idx=scope_end_after,
            ) or title_idx_after
            header_indices_after = [
                idx
                for idx in range(max(udc_indices_after) + 1, min(front_boundary, title_idx_after))
                if idx != title_idx_after
                if (doc.paragraphs[idx - 1].text or "").strip()
            ]
            reordered_front = [doc.paragraphs[idx - 1]._element for idx in header_indices_after]
            reordered_front.append(doc.paragraphs[title_idx_after - 1]._element)
            if reordered_front:
                insert_pos = list(body).index(doc.paragraphs[udc_indices_after[-1] - 1]._element) + 1
                for elem in reordered_front:
                    try:
                        body.remove(elem)
                    except Exception:
                        pass
                for offset, elem in enumerate(reordered_front):
                    body.insert(insert_pos + offset, elem)
                changed = True
                scope_end_after = _front_scope_end_idx(doc)
                title_idx_after = _find_title_idx(doc, expected_title, max_idx=scope_end_after)
                if title_idx_after is not None:
                    changed = _remove_empty_paragraphs_in_range(
                        doc,
                        max(udc_indices_after) + 1,
                        title_idx_after - 1,
                    ) or changed
                    scope_end_after = _front_scope_end_idx(doc)
                    title_idx_after = _find_title_idx(doc, expected_title, max_idx=scope_end_after)
                if title_idx_after is not None:
                    header_indices_after = [
                        idx
                        for idx in range(max(udc_indices_after) + 1, title_idx_after)
                        if (doc.paragraphs[idx - 1].text or "").strip()
                    ]

    front_boundary_after_reorder = _find_front_boundary_idx(doc, max_idx=_front_scope_end_idx(doc))
    if front_boundary_after_reorder and front_boundary_after_reorder > 1:
        changed = _remove_empty_paragraphs_in_range(doc, 1, front_boundary_after_reorder - 1) or changed
        udc_indices_after, title_idx_after, header_indices_after = _recompute_front_indices(doc, expected_title)

    if udc_indices_after:
        udc_last = udc_indices_after[-1]
        udc_style = ""
        try:
            udc_style = doc.paragraphs[udc_last - 1].style.name
        except Exception:
            udc_style = ""
        changed = _ensure_single_blank_after(doc, udc_last, style_name=udc_style or None) or changed
        udc_indices_after, title_idx_after, header_indices_after = _recompute_front_indices(doc, expected_title)

    if header_indices_after:
        header_last = header_indices_after[-1]
        header_style = ""
        try:
            header_style = doc.paragraphs[header_last - 1].style.name
        except Exception:
            header_style = ""
        changed = _ensure_single_blank_after(doc, header_last, style_name=header_style or None) or changed
        udc_indices_after, title_idx_after, header_indices_after = _recompute_front_indices(doc, expected_title)

    if title_idx_after is not None:
        changed = _ensure_single_blank_after(doc, title_idx_after) or changed
    if udc_indices_after:
        changed = _apply_udc_style_to_block(doc, udc_indices_after) or changed
    changed = _ensure_blank_after_top_udc_block(doc) or changed

    if changed:
        doc.save(doc_path)

    result = {
        "changed": changed,
        "title": expected_title,
        "title_detect_mode": title_reason,
        "scope": "first_n_paragraphs",
        "scope_max_paragraphs": FRONT_SCOPE_MAX_PARAGRAPHS,
        "scope_end_idx": scope_end_idx,
        "title_idx_before": title_idx,
        "udc_idx_before": udc_idx,
        "udc_indices_before": udc_indices,
        "header_indices_before": header_indices,
    }
    if logs_dir is not None:
        logs_dir.mkdir(parents=True, exist_ok=True)
        (logs_dir / f"{doc_path.stem}.order_report.json").write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    return result
