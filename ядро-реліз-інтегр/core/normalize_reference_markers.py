from __future__ import annotations

import argparse
import json
import os
from pathlib import Path

from docx import Document

try:
    from core.paragraph_utils import ensure_single_empty_around
except Exception:
    from paragraph_utils import ensure_single_empty_around

try:
    from core.reference_utils import (
        capitalize_first,
        extract_marker_remainder,
        is_cyrillic,
        load_reference_markers,
        reference_heading,
    )
except Exception:
    from reference_utils import (
        capitalize_first,
        extract_marker_remainder,
        is_cyrillic,
        load_reference_markers,
        reference_heading,
    )


def _replace_paragraph_with_heading(para, heading: str, content: str) -> None:
    for run in list(para.runs):
        keep_run = False
        try:
            for node in run._element.iter():
                tag = str(node.tag)
                if tag.endswith("}drawing") or tag.endswith("}pict") or tag.endswith("}object"):
                    keep_run = True
                    break
        except Exception:
            keep_run = False
        if keep_run:
            continue
        run.text = ""
    heading_run = para.add_run(heading)
    heading_run.bold = True
    if content:
        para.add_run(" ")
        para.add_run(content)
    try:
        para.style = "REF-TITLE"
    except Exception:
        pass


def _word_visible() -> bool:
    return os.getenv("WORD_COM_VISIBLE", "").strip().lower() in {"1", "true", "yes", "y"}


def _remove_paragraph(para) -> None:
    element = para._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _trim_trailing_empty_paragraphs(doc: Document) -> int:
    removed = 0
    while doc.paragraphs:
        tail = doc.paragraphs[-1]
        if (tail.text or "").strip():
            break
        _remove_paragraph(tail)
        removed += 1
    return removed


def normalize_references(
    draft_path: Path,
    output_path: Path | None = None,
    *,
    logs_dir: Path | None = None,
    write_logs: bool = True,
) -> Path:
    markers = load_reference_markers()

    doc = Document(draft_path)
    changes = []
    found = 0
    trimmed_tail = 0

    for idx, para in enumerate(doc.paragraphs):
        text = para.text or ""
        remainder, marker = extract_marker_remainder(text, markers)
        if marker is None:
            continue
        found += 1
        lang = "uk" if is_cyrillic(marker or text) else "en"
        heading = reference_heading(lang)
        content = capitalize_first(remainder.strip())
        if content:
            new_text = f"{heading} {content}"
        else:
            new_text = heading
        if (para.text or "") != new_text:
            changes.append(
                {
                    "index": idx + 1,
                    "before": para.text,
                    "after": new_text,
                }
            )
        _replace_paragraph_with_heading(para, heading, content)
        ensure_single_empty_around(para)

    target = output_path or draft_path.with_name(draft_path.stem + "_references.docx")
    trimmed_tail = _trim_trailing_empty_paragraphs(doc)
    doc.save(target)
    _restart_reference_numbering(target)

    if not write_logs:
        return target

    report = {
        "input": str(draft_path),
        "output": str(target),
        "markers_found": found,
        "changed": len(changes),
        "trimmed_trailing_empty": trimmed_tail,
        "changes": changes,
    }
    if logs_dir is None:
        logs_dir = target.parent
    logs_dir.mkdir(parents=True, exist_ok=True)
    report_path = logs_dir / f"{target.stem}.references_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    log_path = logs_dir / f"{target.stem}.references.log"
    lines = []
    for item in changes:
        lines.append(f"[{item['index']}] {item['before']} -> {item['after']}")
    log_path.write_text("\n".join(lines), encoding="utf-8")
    return target


def _restart_reference_numbering(path: Path) -> None:
    try:
        import pythoncom
        import win32com.client
    except Exception:
        return

    headings = {"СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ:", "REFERENCES:"}
    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = bool(_word_visible())
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(path), ReadOnly=False, AddToRecentFiles=False)
        constants = win32com.client.constants
        wdListNoNumbering = 0
        apply_to_forward = getattr(constants, "wdListApplyToThisPointForward", 2)
        default_behavior = getattr(constants, "wdWord10ListBehavior", 2)

        total = int(document.Paragraphs.Count)
        for idx in range(1, total + 1):
            para = document.Paragraphs(idx)
            text = para.Range.Text.strip("\r\x07 ").strip()
            if not text:
                continue
            if text not in headings:
                continue
            for j in range(idx + 1, total + 1):
                next_para = document.Paragraphs(j)
                next_text = next_para.Range.Text.strip("\r\x07 ").strip()
                if not next_text:
                    continue
                if next_text in headings:
                    break
                has_graphics = False
                try:
                    has_graphics = int(next_para.Range.InlineShapes.Count) > 0
                except Exception:
                    has_graphics = False
                if not has_graphics:
                    try:
                        has_graphics = int(next_para.Range.ShapeRange.Count) > 0
                    except Exception:
                        pass
                if has_graphics:
                    continue
                try:
                    list_type = int(next_para.Range.ListFormat.ListType)
                except Exception:
                    list_type = wdListNoNumbering
                if list_type == wdListNoNumbering:
                    break
                try:
                    template = next_para.Range.ListFormat.ListTemplate
                    next_para.Range.ListFormat.ApplyListTemplateWithLevel(
                        template,
                        False,
                        apply_to_forward,
                        default_behavior,
                    )
                except Exception:
                    pass
                break
        document.Save()
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def main() -> None:
    parser = argparse.ArgumentParser(description="Normalize reference markers in draft")
    parser.add_argument("draft", type=Path, help="Draft docx")
    parser.add_argument("--output", type=Path, default=None, help="Output docx")
    args = parser.parse_args()
    if not args.draft.exists():
        raise SystemExit("Draft not found")
    out = normalize_references(args.draft, output_path=args.output)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
