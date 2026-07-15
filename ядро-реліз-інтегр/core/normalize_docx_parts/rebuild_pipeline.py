from __future__ import annotations

from dataclasses import dataclass
import io
import zipfile
import xml.etree.ElementTree as ET

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError

try:
    from core.reference_utils import (
        capitalize_first,
        extract_marker_remainder,
        is_cyrillic,
        load_reference_markers,
        matches_reference_marker,
        reference_heading,
    )
except Exception:
    from reference_utils import (  # type: ignore
        capitalize_first,
        extract_marker_remainder,
        is_cyrillic,
        load_reference_markers,
        matches_reference_marker,
        reference_heading,
    )

try:
    from core.normalize_docx_parts.list_format import next_list_prefix, paragraph_list_info
    from core.normalize_docx_parts.reference_rules import normalize_reference_line_runs
    from core.normalize_docx_parts.table_builder import (
        apply_table_borders,
        format_cell_paragraphs,
        table_layout,
        table_to_text,
    )
    from core.normalize_docx_parts.xml_extract import (
        add_runs,
        image_sources,
        is_supported_image_ext,
        paragraph_chart_count,
        paragraph_image_rids,
        paragraph_runs,
        paragraph_text,
        paragraph_textboxes,
        paragraph_unsupported_drawing_count,
        split_runs_on_breaks,
    )
except Exception:
    from normalize_docx_parts.list_format import next_list_prefix, paragraph_list_info  # type: ignore
    from normalize_docx_parts.reference_rules import normalize_reference_line_runs  # type: ignore
    from normalize_docx_parts.table_builder import (  # type: ignore
        apply_table_borders,
        format_cell_paragraphs,
        table_layout,
        table_to_text,
    )
    from normalize_docx_parts.xml_extract import (  # type: ignore
        add_runs,
        image_sources,
        is_supported_image_ext,
        paragraph_chart_count,
        paragraph_image_rids,
        paragraph_runs,
        paragraph_text,
        paragraph_textboxes,
        paragraph_unsupported_drawing_count,
        split_runs_on_breaks,
    )


@dataclass(frozen=True)
class RebuildResult:
    chart_placeholder_idx: int
    image_placeholder_positions: list[int]
    math_placeholder_count: int
    unsupported_drawing_count: int


def _add_empty_paragraph(out) -> None:
    out.add_paragraph("")


def _add_unsupported_placeholder(out, idx: int) -> None:
    out.add_paragraph(f"<<UNSUPPORTED_DRAWING_{idx}>>")


def _strip_line_leading_ws(
    runs: list[tuple[str, bool, bool, bool, bool]]
) -> list[tuple[str, bool, bool, bool, bool]]:
    # Some source docs emulate paragraph tabs via leading spaces in first run.
    # Remove only line-leading whitespace, keep internal spacing intact.
    out: list[tuple[str, bool, bool, bool, bool]] = []
    stripped = False
    for text, bold, italic, superscript, subscript in runs:
        if not stripped:
            cleaned = text.lstrip(" \t\u00A0")
            if cleaned:
                out.append((cleaned, bold, italic, superscript, subscript))
                stripped = True
            else:
                continue
        else:
            out.append((text, bold, italic, superscript, subscript))
    return out


def rebuild_document_from_body(
    *,
    body: ET.Element,
    out,
    zf: zipfile.ZipFile,
    rels: dict[str, str],
    numbering: dict[tuple[str, int], str],
    ns: dict[str, str],
    refer_style: str,
    table_text_style: str,
) -> RebuildResult:
    in_references = False
    counters: dict[tuple[str, int], int] = {}
    chart_placeholder_idx = 0
    image_placeholder_idx = 0
    image_seq_idx = 0
    image_placeholder_positions: list[int] = []
    math_state = {"count": 0}
    unsupported_drawing_count = 0
    reference_markers = load_reference_markers()

    for child in list(body):
        tag = child.tag
        if tag == f"{{{ns['w']}}}p":
            if not in_references and reference_markers:
                raw_text = paragraph_text(child, ns)
                if matches_reference_marker(raw_text, reference_markers):
                    remainder, marker = extract_marker_remainder(raw_text, reference_markers)
                    lang = "uk" if is_cyrillic(marker or raw_text) else "en"
                    heading = reference_heading(lang)
                    if out.paragraphs and out.paragraphs[-1].text.strip():
                        _add_empty_paragraph(out)
                    heading_para = out.add_paragraph()
                    try:
                        heading_para.style = "REF-TITLE"
                    except Exception:
                        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = heading_para.add_run(heading)
                    run.bold = True
                    remainder = capitalize_first(remainder.strip())
                    if remainder:
                        heading_para.add_run(" ")
                        heading_para.add_run(remainder)
                    _add_empty_paragraph(out)
                    in_references = True
                    continue

            runs = paragraph_runs(child, ns, math_state)
            if runs:
                list_info = paragraph_list_info(child, numbering, ns)
                lines = split_runs_on_breaks(runs)
                if in_references:
                    para = out.add_paragraph()
                    try:
                        para.style = refer_style
                    except Exception:
                        pass
                    for idx, line in enumerate(lines):
                        line = _strip_line_leading_ws(line)
                        line = normalize_reference_line_runs(line)
                        if idx > 0:
                            para.add_run(" ")
                        add_runs(para, line)
                elif list_info and not in_references:
                    indent, prefix = next_list_prefix(list_info, counters)
                    for line in lines:
                        line = _strip_line_leading_ws(line)
                        para = out.add_paragraph()
                        para.add_run(f"{indent}{prefix}")
                        add_runs(para, line)
                else:
                    dash_prefix = "- " if len(lines) > 1 and not in_references else ""
                    for line in lines:
                        line = _strip_line_leading_ws(line)
                        para = out.add_paragraph()
                        if in_references:
                            try:
                                para.style = refer_style
                            except Exception:
                                pass
                        if dash_prefix:
                            para.add_run(dash_prefix)
                        add_runs(para, line)

            rids = paragraph_image_rids(child, ns)
            if rids:
                for blob, ext in image_sources(zf, rels, rids):
                    image_seq_idx += 1
                    if not is_supported_image_ext(ext):
                        image_placeholder_idx += 1
                        image_placeholder_positions.append(image_seq_idx)
                        out.add_paragraph(f"<<IMG_PLACEHOLDER_{image_placeholder_idx}>>")
                        continue
                    try:
                        out.add_picture(io.BytesIO(blob))
                    except UnrecognizedImageError:
                        image_placeholder_idx += 1
                        image_placeholder_positions.append(image_seq_idx)
                        out.add_paragraph(f"<<IMG_PLACEHOLDER_{image_placeholder_idx}>>")

            chart_count = paragraph_chart_count(child, ns)
            for _ in range(chart_count):
                chart_placeholder_idx += 1
                out.add_paragraph(f"<<CHART_PLACEHOLDER_{chart_placeholder_idx}>>")

            unsupported_count = paragraph_unsupported_drawing_count(child, ns)
            for _ in range(unsupported_count):
                unsupported_drawing_count += 1
                _add_unsupported_placeholder(out, unsupported_drawing_count)

            for box_text in paragraph_textboxes(child, ns):
                for line in box_text.split("\n"):
                    out.add_paragraph(line)

        elif tag == f"{{{ns['w']}}}tbl":
            rows = child.findall("w:tr", ns)
            if not rows:
                continue
            logical_cols, rows_layout, table_merges = table_layout(child, ns)
            table = out.add_table(rows=len(rows), cols=logical_cols)
            apply_table_borders(table)

            for r_idx, row_layout in enumerate(rows_layout):
                for tc, c_idx, _span in row_layout:
                    paragraphs = tc.findall("w:p", ns)
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ""
                    has_any = False

                    for p in paragraphs:
                        runs = paragraph_runs(p, ns, math_state)
                        if not runs:
                            continue
                        lines = split_runs_on_breaks(runs)
                        list_info = paragraph_list_info(p, numbering, ns)
                        if in_references:
                            para = cell.paragraphs[-1] if not has_any else cell.add_paragraph()
                            try:
                                para.style = refer_style
                            except Exception:
                                pass
                            for idx, line in enumerate(lines):
                                line = _strip_line_leading_ws(line)
                                line = normalize_reference_line_runs(line)
                                if idx > 0:
                                    para.add_run(" ")
                                add_runs(para, line)
                            has_any = True
                        elif list_info and not in_references:
                            indent, prefix = next_list_prefix(list_info, counters)
                            for line in lines:
                                line = _strip_line_leading_ws(line)
                                para = cell.paragraphs[-1] if not has_any else cell.add_paragraph()
                                para.add_run(f"{indent}{prefix}")
                                add_runs(para, line)
                                has_any = True
                        else:
                            dash_prefix = "- " if len(lines) > 1 and not in_references else ""
                            for line in lines:
                                line = _strip_line_leading_ws(line)
                                para = cell.paragraphs[-1] if not has_any else cell.add_paragraph()
                                if in_references:
                                    try:
                                        para.style = refer_style
                                    except Exception:
                                        pass
                                if dash_prefix:
                                    para.add_run(dash_prefix)
                                add_runs(para, line)
                                has_any = True

                    nested_text = table_to_text(tc, numbering, counters, ns)
                    if nested_text:
                        para = cell.paragraphs[-1] if not has_any else cell.add_paragraph()
                        para.add_run(nested_text)
                        has_any = True
                    format_cell_paragraphs(cell, style_name=table_text_style)

                    for p in paragraphs:
                        rids = paragraph_image_rids(p, ns)
                        for blob, ext in image_sources(zf, rels, rids):
                            image_seq_idx += 1
                            if not is_supported_image_ext(ext):
                                image_placeholder_idx += 1
                                image_placeholder_positions.append(image_seq_idx)
                                cell.add_paragraph(f"<<IMG_PLACEHOLDER_{image_placeholder_idx}>>")
                                continue
                            try:
                                run = cell.paragraphs[-1].add_run()
                                run.add_picture(io.BytesIO(blob))
                            except UnrecognizedImageError:
                                image_placeholder_idx += 1
                                image_placeholder_positions.append(image_seq_idx)
                                cell.add_paragraph(f"<<IMG_PLACEHOLDER_{image_placeholder_idx}>>")

                        chart_count = paragraph_chart_count(p, ns)
                        for _ in range(chart_count):
                            chart_placeholder_idx += 1
                            cell.add_paragraph(f"<<CHART_PLACEHOLDER_{chart_placeholder_idx}>>")

                        unsupported_count = paragraph_unsupported_drawing_count(p, ns)
                        for _ in range(unsupported_count):
                            unsupported_drawing_count += 1
                            cell.add_paragraph(f"<<UNSUPPORTED_DRAWING_{unsupported_drawing_count}>>")

                        for box_text in paragraph_textboxes(p, ns):
                            cell.add_paragraph(box_text)
                    format_cell_paragraphs(cell, style_name=table_text_style)

            for start_row, start_col, end_row, end_col in table_merges:
                try:
                    table.cell(start_row, start_col).merge(table.cell(end_row, end_col))
                except (IndexError, ValueError):
                    continue

    return RebuildResult(
        chart_placeholder_idx=chart_placeholder_idx,
        image_placeholder_positions=image_placeholder_positions,
        math_placeholder_count=math_state["count"],
        unsupported_drawing_count=unsupported_drawing_count,
    )
