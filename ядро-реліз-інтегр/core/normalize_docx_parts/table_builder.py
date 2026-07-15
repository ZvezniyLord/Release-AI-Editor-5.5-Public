from __future__ import annotations

from pathlib import Path
import tempfile
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .list_format import next_list_prefix, paragraph_list_info
from .xml_extract import paragraph_text


def table_layout(
    tbl: ET.Element,
    ns: dict[str, str],
) -> tuple[int, list[list[tuple[ET.Element, int, int]]], list[tuple[int, int, int, int]]]:
    grid_cols = len(tbl.findall("w:tblGrid/w:gridCol", ns))
    rows_layout: list[list[tuple[ET.Element, int, int]]] = []
    horizontal_merges: list[tuple[int, int, int, int]] = []
    vertical_starts: dict[tuple[int, int], int] = {}
    vertical_merges: list[tuple[int, int, int, int]] = []
    max_logical_cols = 0

    for row_idx, tr in enumerate(tbl.findall("w:tr", ns)):
        logical_col = 0
        row_layout: list[tuple[ET.Element, int, int]] = []
        for tc in tr.findall("w:tc", ns):
            span_node = tc.find("w:tcPr/w:gridSpan", ns)
            try:
                span = max(1, int(span_node.get(f"{{{ns['w']}}}val"))) if span_node is not None else 1
            except (TypeError, ValueError):
                span = 1
            row_layout.append((tc, logical_col, span))

            merge_node = tc.find("w:tcPr/w:vMerge", ns)
            if merge_node is not None:
                merge_value = merge_node.get(f"{{{ns['w']}}}val") or "continue"
                key = (logical_col, span)
                if merge_value == "restart":
                    vertical_starts[key] = row_idx
                elif key in vertical_starts:
                    vertical_merges.append(
                        (vertical_starts[key], logical_col, row_idx, logical_col + span - 1)
                    )
            if span > 1:
                horizontal_merges.append(
                    (row_idx, logical_col, row_idx, logical_col + span - 1)
                )
            logical_col += span
        max_logical_cols = max(max_logical_cols, logical_col)
        rows_layout.append(row_layout)

    return max(grid_cols, max_logical_cols, 1), rows_layout, horizontal_merges + vertical_merges


def apply_table_borders(table) -> None:
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    tbl = table._element
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is not None:
        tbl_pr.remove(tbl_borders)

    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        tbl_borders.append(elem)
    tbl_pr.append(tbl_borders)


def set_cell_margins_zero(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:tcMar"):
            tc_pr.remove(child)
    tc_mar = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), "0")
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def format_cell_paragraphs(cell, *, style_name: str | None = None) -> None:
    set_cell_margins_zero(cell)
    for paragraph in cell.paragraphs:
        if style_name:
            try:
                paragraph.style = style_name
            except Exception:
                pass
        fmt = paragraph.paragraph_format
        fmt.line_spacing = 1.15
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = Pt(0)
        fmt.right_indent = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"


def table_to_text(
    tbl: ET.Element,
    numbering: dict[tuple[str, int], str],
    counters: dict[tuple[str, int], int],
    ns: dict[str, str],
) -> str:
    rows = tbl.findall("w:tr", ns)
    out_rows: list[str] = []
    for tr in rows:
        cells = tr.findall("w:tc", ns)
        row_cells: list[str] = []
        for tc in cells:
            parts: list[str] = []
            for p in tc.findall("w:p", ns):
                t = paragraph_text(p, ns)
                if not t:
                    continue
                list_info = paragraph_list_info(p, numbering, ns)
                if list_info:
                    indent, prefix = next_list_prefix(list_info, counters)
                    parts.append(f"{indent}{prefix}{t}")
                else:
                    parts.append(t)
            # nested tables -> flatten into text
            for nested in tc.findall(".//w:tbl", ns):
                nested_rows: list[str] = []
                for ntr in nested.findall("w:tr", ns):
                    n_cells = ntr.findall("w:tc", ns)
                    n_row: list[str] = []
                    for ntc in n_cells:
                        n_texts: list[str] = []
                        for np in ntc.findall("w:p", ns):
                            t = paragraph_text(np, ns)
                            if not t:
                                continue
                            list_info = paragraph_list_info(np, numbering, ns)
                            if list_info:
                                indent, prefix = next_list_prefix(list_info, counters)
                                n_texts.append(f"{indent}{prefix}{t}")
                            else:
                                n_texts.append(t)
                        n_row.append(" ".join(n_texts))
                    if n_row:
                        nested_rows.append(" | ".join(n_row))
                if nested_rows:
                    parts.append(" / ".join(nested_rows))
            cell_text = " ".join(parts).strip()
            row_cells.append(cell_text)
        if row_cells:
            out_rows.append(" | ".join(row_cells))
    return "\n".join(out_rows).strip()


def ensure_table_text_style(doc: Document) -> str:
    name = "TABLETEXT"
    try:
        style = doc.styles[name]
    except Exception:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        try:
            style.base_style = doc.styles["Normal"]
        except Exception:
            pass

    fmt = style.paragraph_format
    fmt.line_spacing = 1.15
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    return name


def ensure_refer_style(doc: Document) -> str:
    name = "REFER"
    try:
        style = doc.styles[name]
    except Exception:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        try:
            style.base_style = doc.styles["Normal"]
        except Exception:
            pass
    return name


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def load_template_document(template_path: Path) -> Document:
    if not template_path.exists():
        return Document()
    if template_path.suffix.lower() != ".dotx":
        return Document(template_path)
    temp_path = None
    try:
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        temp_path = Path(temp.name)
        temp.close()
        with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(temp_path, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                    )
                zout.writestr(item, data)
        return Document(temp_path)
    finally:
        if temp_path and temp_path.exists():
            try:
                temp_path.unlink()
            except Exception:
                pass
