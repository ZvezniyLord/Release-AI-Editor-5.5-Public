from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


MAX_HEADER_PARAS = 35
MAX_LINE_LENGTH = 180
MIN_LINE_LEN = 3
TITLE_PREFIX_RE = re.compile(
    r"^\s*(?:тези|тезисы|theses|article|стаття)\s*[:\-–—]?\s*",
    flags=re.IGNORECASE,
)


def _cp1251_mojibake(text: str) -> str:
    return text.encode("utf-8").decode("cp1251")


def _strip_outer_quotes(text: str) -> str:
    value = (text or "").strip()
    if len(value) < 2:
        return value
    quote_pairs = [
        ('"', '"'),
        ("'", "'"),
        ("“", "”"),
        ("„", "“"),
        ("«", "»"),
        ("‹", "›"),
        (_cp1251_mojibake("“"), _cp1251_mojibake("”")),
        (_cp1251_mojibake("«"), _cp1251_mojibake("»")),
    ]
    changed = True
    while changed and len(value) >= 2:
        changed = False
        for left, right in quote_pairs:
            if value.startswith(left) and value.endswith(right):
                inner = value[len(left): len(value) - len(right)].strip()
                if inner:
                    value = inner
                    changed = True
                break
    return value


def _normalize_title_source(text: str) -> str:
    value = (text or "").strip()
    if not value:
        return ""
    value = TITLE_PREFIX_RE.sub("", value)
    value = _strip_outer_quotes(value)
    return value.strip()


def _normalize_key(text: str) -> str:
    text = _normalize_title_source(text).casefold()
    text = re.sub(r"[^a-zа-яіїєґ0-9]+", " ", text, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", text).strip()


def _normalize_key_compact(text: str) -> str:
    text = _normalize_title_source(text).casefold()
    return re.sub(r"[^a-zа-яіїєґ0-9]+", "", text, flags=re.IGNORECASE)


def _load_titles_from_matches(draft_path: Path) -> tuple[list[str], dict]:
    matches_path = draft_path.parent / "matches.json"
    if not matches_path.exists():
        return [], {"total": 0, "title_match": 0, "free_listener": 0, "missing": 0, "other": 0}
    try:
        data = json.loads(matches_path.read_text(encoding="utf-8"))
    except Exception:
        return [], {"total": 0, "title_match": 0, "free_listener": 0, "missing": 0, "other": 0}
    if not isinstance(data, list):
        return [], {"total": 0, "title_match": 0, "free_listener": 0, "missing": 0, "other": 0}
    titles: list[str] = []
    stats = {"total": len(data), "title_match": 0, "free_listener": 0, "missing": 0, "other": 0}
    for item in data:
        if not isinstance(item, dict):
            continue
        method = (item.get("match_method") or "").strip()
        if method == "title_match":
            stats["title_match"] += 1
        elif method == "free_listener":
            stats["free_listener"] += 1
        elif method == "missing":
            stats["missing"] += 1
        else:
            stats["other"] += 1
        if method != "title_match":
            continue
        title = (item.get("title") or "").strip()
        if title:
            titles.append(title)
    return titles, stats


def _load_matches_by_title(draft_path: Path) -> dict[str, dict]:
    matches_path = draft_path.parent / "matches.json"
    if not matches_path.exists():
        return {}
    try:
        data = json.loads(matches_path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    if not isinstance(data, list):
        return {}
    result: dict[str, dict] = {}
    for item in data:
        if not isinstance(item, dict):
            continue
        title = (item.get("title") or "").strip()
        if title and title not in result:
            result[title] = item
    return result


def _remove_paragraph(paragraph) -> None:
    try:
        p = paragraph._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)
    except Exception:
        pass


def _merge_title_paragraphs(
    doc: Document,
    title_keys: dict[str, str],
    title_keys_compact: set[str],
) -> tuple[set[int], dict[str, int]]:
    """
    Зливає 2-6 абзаців, якщо їхній сумарний текст == назві зі списку.
    Повертає індекси (1-based) абзаців, де стоїть назва після злиття.
    """
    if not title_keys:
        return set(), {}
    paras = doc.paragraphs
    keys = set(title_keys.keys())
    hits: set[int] = set()
    merged_counts: dict[str, int] = {"2": 0, "3": 0, "4": 0, "5": 0, "6": 0}

    idx = 0
    while idx < len(paras):
        t1 = paras[idx].text.strip()
        k1 = _normalize_key(t1)
        k1c = _normalize_key_compact(t1)
        if (k1 and k1 in keys) or (k1c and k1c in title_keys_compact):
            hits.add(idx + 1)
            idx += 1
            continue

        if t1:
            for span in range(2, 7):
                if idx + span - 1 >= len(paras):
                    break
                parts = []
                ok = True
                for j in range(span):
                    t = paras[idx + j].text.strip()
                    if not t:
                        ok = False
                        break
                    parts.append(t)
                if not ok:
                    continue
                joined = " ".join(parts)
                key = _normalize_key(joined)
                keyc = _normalize_key_compact(joined)
                if key in keys or keyc in title_keys_compact:
                    paras[idx].text = " ".join(parts)
                    for j in range(span - 1, 0, -1):
                        _remove_paragraph(paras[idx + j])
                    hits.add(idx + 1)
                    merged_counts[str(span)] += 1
                    paras = doc.paragraphs
                    break
            if (idx + 1) in hits:
                continue

        idx += 1
    return hits, merged_counts


def _looks_like_title(text: str) -> bool:
    letters = [c for c in text if c.isalpha()]
    if len(letters) < 12 or len(text) > 220:
        return False
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    return upper_ratio >= 0.55 or len(text) > 70


def _collect_title_candidates(doc: Document) -> dict[str, list[int]]:
    candidates = {"title": []}
    since_break = 0

    def has_page_break(p) -> bool:
        for run in p.runs:
            el = run._element
            if el is None:
                continue
            for node in el.iter():
                if node.tag.endswith("}br") and node.get(qn("w:type")) == "page":
                    return True
                if node.tag.endswith("}lastRenderedPageBreak"):
                    return True
        return False

    for idx, para in enumerate(doc.paragraphs, start=1):
        raw_text = para.text or ""
        if "\x0c" in raw_text or has_page_break(para):
            since_break = 0
            continue
        since_break += 1
        text = raw_text.strip()
        if not text:
            continue
        if len(text) < MIN_LINE_LEN or len(text) > MAX_LINE_LENGTH:
            continue
        if since_break > MAX_HEADER_PARAS:
            continue
        is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        if _looks_like_title(text) and (is_centered or text.isupper()):
            candidates["title"].append(idx)
    return candidates
