# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
import json
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import fitz

ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = Path(__file__).resolve().parents[3]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from audit import EXPECTED_COLUMN_WIDTHS_TWIPS, audit_toc_docx, write_audit_json
from manifest_reader import load_free_listeners
from outline_parser import build_sections, parse_outline
from table_builder import build_rows

FIXTURE_DIR = REPO_ROOT / "fixtures" / "synthetic" / "toc"
ARTIFACT_DIR = REPO_ROOT / "artifacts" / "toc_vertical_slice"

SOURCE_DOCX = FIXTURE_DIR / "TOC_VERTICAL_SLICE_INPUT.docx"
MANIFEST_JSON = FIXTURE_DIR / "manifest.json"
MATCHES_JSON = FIXTURE_DIR / "matches.json"
SECTIONS_JSON = FIXTURE_DIR / "sections.json"

OUTPUT_DOCX = ARTIFACT_DIR / "TOC_VERTICAL_SLICE.docx"
OUTPUT_PDF = ARTIFACT_DIR / "TOC_VERTICAL_SLICE.pdf"
CONTACT_SHEET = ARTIFACT_DIR / "contact_sheet.png"
AUDIT_JSON = ARTIFACT_DIR / "TOC_AUDIT.json"

SYNTHETIC_SECTIONS = [
    (
        "Synthetic Section One",
        [
            ("Marta Testova; Danylo Pryklad", "Synthetic Article on Local Editorial Workflow"),
            ("Oksana Vyhadana", "Synthetic Article on Table Contract Verification"),
        ],
    ),
    (
        "Synthetic Section Two",
        [
            ("Ivan Navchalnyi", "Synthetic Article on Fail Closed TOC Audits"),
        ],
    ),
]

FREE_LISTENER_HEADER = "SYNTHETIC FREE LISTENERS"
FREE_LISTENER_NAME = "Free Listener Test Participant"


def _ensure_paragraph_style(document: Document, name: str, *, size: int = 11, bold: bool = False) -> None:
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = bold


def _create_source_fixture(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.author = "Synthetic Fixture"
    document.core_properties.last_modified_by = "Synthetic Fixture"
    document.core_properties.title = "TOC Vertical Slice Synthetic Input"
    for style_name in ("SECTION", "AUTOR", "Назва1"):
        _ensure_paragraph_style(document, style_name)

    service = document.add_paragraph("Synthetic Service Page")
    service.style = "Title"
    document.add_paragraph("анкета: synthetic noise fragment that must not enter TOC")

    for section_name, articles in SYNTHETIC_SECTIONS:
        section = document.add_paragraph(section_name)
        section.style = "SECTION"
        for authors, title in articles:
            author_p = document.add_paragraph(authors)
            author_p.style = "AUTOR"
            title_p = document.add_paragraph(title)
            title_p.style = "Назва1"

    document.save(path)


def _create_manifest_fixtures() -> None:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    MATCHES_JSON.write_text(
        json.dumps(
            {"matches": [{"match_method": "free_listener", "authors": [FREE_LISTENER_NAME]}]},
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    SECTIONS_JSON.write_text(
        json.dumps([{"block_number": 0, "section_en": FREE_LISTENER_HEADER}], ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    MANIFEST_JSON.write_text(
        json.dumps(
            {"matches_json_path": "matches.json", "sections_path": "sections.json"},
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )


def _set_table_fixed_layout(table) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for grid_col, width in zip(tbl.tblGrid.gridCol_lst, EXPECTED_COLUMN_WIDTHS_TWIPS):
        grid_col.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, EXPECTED_COLUMN_WIDTHS_TWIPS):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def _clear_cell(cell) -> None:
    cell.text = ""


def _create_toc_docx(path: Path, rows) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.author = "Synthetic Fixture"
    document.core_properties.last_modified_by = "Synthetic Fixture"
    document.core_properties.title = "TOC Vertical Slice Synthetic Output"
    for style_name, size, bold in (
        ("Tab_SEC", 11, True),
        ("Tab_PIP", 10, False),
        ("Tab_Taitl", 10, False),
    ):
        _ensure_paragraph_style(document, style_name, size=size, bold=bold)

    table = document.add_table(rows=len(rows), cols=3)
    _set_table_fixed_layout(table)

    for table_row, row in zip(table.rows, rows):
        for cell in table_row.cells:
            _clear_cell(cell)
        center = table_row.cells[1]
        if row.kind == "section":
            p = center.paragraphs[0]
            p.text = row.section
            p.style = "Tab_SEC"
        elif row.kind == "free_listeners_header":
            p = center.paragraphs[0]
            p.text = row.section
            p.style = "Tab_SEC"
        elif row.kind == "item":
            p = center.paragraphs[0]
            p.text = row.authors
            p.style = "Tab_PIP"
            title = center.add_paragraph(row.title)
            title.style = "Tab_Taitl"
        elif row.kind == "free_listeners":
            p = center.paragraphs[0]
            p.text = row.free_listeners
            p.style = "Tab_PIP"
    document.save(path)


def _row_display_text(row) -> list[str]:
    if row.kind == "item":
        return [row.authors, row.title]
    if row.kind in {"section", "free_listeners_header"}:
        return [row.section]
    return [row.free_listeners]


def _create_pdf(path: Path, rows) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    left = 56
    top = 72
    widths = [w / 20 for w in EXPECTED_COLUMN_WIDTHS_TWIPS]
    row_heights = [34 if row.kind in {"section", "free_listeners_header"} else 52 for row in rows]
    y = top
    for row, height in zip(rows, row_heights):
        x = left
        for width in widths:
            page.draw_rect(fitz.Rect(x, y, x + width, y + height), color=(0, 0, 0), width=0.5)
            x += width
        text_rect = fitz.Rect(left + widths[0] + 6, y + 6, left + widths[0] + widths[1] - 6, y + height - 4)
        text = "\n".join(_row_display_text(row))
        fontsize = 8.8 if row.kind == "item" else 9.5
        page.insert_textbox(text_rect, text, fontsize=fontsize, fontname="helv", color=(0, 0, 0))
        y += height
    doc.save(path)
    doc.close()


def _create_contact_sheet(pdf_path: Path, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(pdf_path)
    page = doc.load_page(0)
    pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
    pix.save(output_path)
    doc.close()


def generate() -> dict:
    _create_source_fixture(SOURCE_DOCX)
    _create_manifest_fixtures()

    free_listener_header, free_listeners = load_free_listeners(MANIFEST_JSON)
    sections = build_sections(parse_outline(SOURCE_DOCX))
    rows = build_rows(sections, free_listeners, free_listener_header=free_listener_header)

    _create_toc_docx(OUTPUT_DOCX, rows)
    _create_pdf(OUTPUT_PDF, rows)
    _create_contact_sheet(OUTPUT_PDF, CONTACT_SHEET)

    audit = audit_toc_docx(OUTPUT_DOCX)
    audit["docx"] = OUTPUT_DOCX.relative_to(REPO_ROOT).as_posix()
    write_audit_json(audit, AUDIT_JSON)
    return audit


if __name__ == "__main__":
    result = generate()
    print(json.dumps(result, ensure_ascii=False, indent=2))
    if result.get("status") != "PASS":
        raise SystemExit(1)
