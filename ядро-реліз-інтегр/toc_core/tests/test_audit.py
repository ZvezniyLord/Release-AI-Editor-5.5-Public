# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
import sys
import pytest

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from audit import EXPECTED_COLUMN_WIDTHS_TWIPS, TOC_VISUAL_CONTRACT_INVALID, audit_toc_docx


pytestmark = pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx not available")


def _ensure_style(document: Document, name: str) -> None:
    try:
        document.styles[name]
    except KeyError:
        document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _set_grid_widths(table, widths: list[int]) -> None:
    table.autofit = False
    layout = table._tbl.tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table._tbl.tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))


def _write_contract_docx(path: Path, widths: list[int]) -> None:
    document = Document()
    for style in ("Tab_SEC", "Tab_PIP", "Tab_Taitl"):
        _ensure_style(document, style)
    table = document.add_table(rows=4, cols=3)
    _set_grid_widths(table, widths)

    table.rows[0].cells[1].paragraphs[0].text = "Synthetic Section"
    table.rows[0].cells[1].paragraphs[0].style = "Tab_SEC"
    table.rows[1].cells[1].paragraphs[0].text = "Synthetic Author"
    table.rows[1].cells[1].paragraphs[0].style = "Tab_PIP"
    title = table.rows[1].cells[1].add_paragraph("Synthetic Title")
    title.style = "Tab_Taitl"
    table.rows[2].cells[1].paragraphs[0].text = "Synthetic Free Listener Header"
    table.rows[2].cells[1].paragraphs[0].style = "Tab_SEC"
    table.rows[3].cells[1].paragraphs[0].text = "Synthetic Listener"
    table.rows[3].cells[1].paragraphs[0].style = "Tab_PIP"
    document.save(path)


def test_toc_audit_passes_valid_contract(tmp_path: Path):
    docx_path = tmp_path / "valid.docx"
    _write_contract_docx(docx_path, EXPECTED_COLUMN_WIDTHS_TWIPS)

    result = audit_toc_docx(
        docx_path,
        expected_journal_sections=1,
        expected_articles=1,
        expected_free_listener_rows=1,
    )

    assert result["status"] == "PASS"
    assert result["errors"] == []


def test_toc_audit_fails_invalid_widths(tmp_path: Path):
    docx_path = tmp_path / "invalid.docx"
    _write_contract_docx(docx_path, [700, 8000, 900])

    result = audit_toc_docx(
        docx_path,
        expected_journal_sections=1,
        expected_articles=1,
        expected_free_listener_rows=1,
    )

    assert result["status"] == "FAIL"
    assert result["errors"][0]["code"] == TOC_VISUAL_CONTRACT_INVALID
