# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
import sys
import pytest

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.style import WD_STYLE_TYPE
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from outline_parser import TocInputError, parse_outline, build_sections


pytestmark = pytest.mark.skipif(not _DOCX_AVAILABLE, reason="python-docx not available")


def _write_doc(tmp_path: Path, paragraphs: list[tuple[str, str]]) -> Path:
    doc = Document()
    for text, style in paragraphs:
        p = doc.add_paragraph(text)
        if style:
            p.style = style
    path = tmp_path / "sample.docx"
    doc.save(path)
    return path


def test_outline_single_item(tmp_path: Path):
    doc_path = _write_doc(tmp_path, [
        ("Section 1", "Heading 1"),
        ("Author A", "Heading 2"),
        ("Title A", "Heading 3"),
    ])
    outline = parse_outline(doc_path)
    sections = build_sections(outline)
    assert len(sections) == 1
    assert sections[0].name == "Section 1"
    assert len(sections[0].items) == 1
    assert sections[0].items[0].authors == "Author A"
    assert sections[0].items[0].title == "Title A"


def test_outline_two_titles_same_author(tmp_path: Path):
    doc_path = _write_doc(tmp_path, [
        ("Section 1", "Heading 1"),
        ("Author A", "Heading 2"),
        ("Title A", "Heading 3"),
        ("Title B", "Heading 3"),
    ])
    outline = parse_outline(doc_path)
    with pytest.raises(TocInputError, match="title without author"):
        build_sections(outline)


def test_outline_missing_authors(tmp_path: Path):
    doc_path = _write_doc(tmp_path, [
        ("Section 1", "Heading 1"),
        ("Title A", "Heading 3"),
    ])
    outline = parse_outline(doc_path)
    with pytest.raises(TocInputError, match="title without author"):
        build_sections(outline)


def test_outline_ignores_empty(tmp_path: Path):
    doc_path = _write_doc(tmp_path, [
        ("", "Heading 1"),
        ("   ", "Heading 2"),
        ("Section 1", "Heading 1"),
        ("Author A", "Heading 2"),
        ("Title A", "Heading 3"),
    ])
    outline = parse_outline(doc_path)
    sections = build_sections(outline)
    assert len(sections) == 1


def test_outline_heading2_with_level3_outline(tmp_path: Path):
    doc = Document()
    p1 = doc.add_paragraph("Section 1")
    p1.style = "Heading 1"
    p2 = doc.add_paragraph("Author A")
    p2.style = "Heading 2"
    p3 = doc.add_paragraph("Title A")
    p3.style = "Heading 2"
    p3._p.get_or_add_pPr().get_or_add_outlineLvl().set(qn("w:val"), "2")
    path = tmp_path / "sample2.docx"
    doc.save(path)

    outline = parse_outline(path)
    sections = build_sections(outline)
    assert len(sections) == 1
    assert len(sections[0].items) == 1
    assert sections[0].items[0].authors == "Author A"
    assert sections[0].items[0].title == "Title A"


def test_outline_multiple_authors_before_title(tmp_path: Path):
    doc_path = _write_doc(tmp_path, [
        ("Section 1", "Heading 1"),
        ("Author A", "Heading 2"),
        ("Author B", "Heading 2"),
        ("Title A", "Heading 3"),
    ])
    outline = parse_outline(doc_path)
    sections = build_sections(outline)
    assert sections[0].items[0].authors == "Author A, Author B"


def test_custom_journal_styles_ignore_service_page_and_noise(tmp_path: Path):
    doc = Document()
    for style_name in ("SECTION", "AUTOR", "Назва1"):
        if style_name not in [style.name for style in doc.styles]:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    service = doc.add_paragraph("Synthetic service page")
    service.style = "Title"
    doc.add_paragraph("анкета: synthetic noise fragment")
    section = doc.add_paragraph("Synthetic Section One")
    section.style = "SECTION"
    author = doc.add_paragraph("Marta Testova")
    author.style = "AUTOR"
    title = doc.add_paragraph("Synthetic Article Title")
    title.style = "Назва1"

    path = tmp_path / "synthetic.docx"
    doc.save(path)

    outline = parse_outline(path)
    assert [item.text for item in outline] == [
        "Synthetic Section One",
        "Marta Testova",
        "Synthetic Article Title",
    ]
    assert [item.level for item in outline] == [1, 2, 3]
    sections = build_sections(outline)
    assert sections[0].items[0].authors == "Marta Testova"
    assert sections[0].items[0].title == "Synthetic Article Title"


def test_author_without_title_fails(tmp_path: Path):
    doc_path = _write_doc(tmp_path, [
        ("Section 1", "Heading 1"),
        ("Author A", "Heading 2"),
    ])
    outline = parse_outline(doc_path)
    with pytest.raises(TocInputError, match="author without title"):
        build_sections(outline)
